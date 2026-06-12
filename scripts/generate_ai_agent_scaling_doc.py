from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "AI客服高并发治理面试方案.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_run(paragraph, text, bold=False, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note(doc, title, body, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [6.5])
    set_cell_margins(table, top=120, bottom=120, start=180, end=180)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    add_run(p, title + "：", bold=True, color="1F3A5F")
    p.add_run(body)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, widths)
    set_cell_margins(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header, bold=True, color="1F3A5F")
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.add_run(text)
    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_doc():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI 客服高并发治理面试方案")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("以 Nexus 电商智能客服 Agent 为例：从 MVP 到可承载大规模并发的演进路径")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string("555555")

    add_note(
        doc,
        "一句话总览",
        "几千人同时使用 AI 客服时，不能让所有请求同步直连大模型。正确做法是：入口限流、超时熔断、任务排队、Worker 池控制 LLM 并发、SSE/WebSocket 流式返回、缓存高频问题，并通过压测和监控不断校准容量。",
        fill="E8EEF5",
    )

    doc.add_heading("1. 面试中先怎么定性问题", level=1)
    add_para(
        doc,
        "面试回答不要从“加机器”开始，而要先指出问题本质：AI 客服不是普通 CRUD 接口，它的慢点在外部 LLM、RAG 检索和上下文工具调用。一次 Agent 回复可能占用后端线程数秒甚至十几秒；如果几千人同时发送消息，Tomcat 线程、数据库连接池、Redis、pgvector 和外部模型 API 都可能成为瓶颈。",
    )
    add_para(
        doc,
        "一个高质量回答应该体现三件事：第一，知道瓶颈在哪里；第二，知道如何保护系统不被打穿；第三，知道如何通过异步化、缓存和水平扩容提升容量。",
    )

    add_matrix(
        doc,
        ["概念", "含义", "面试表达"],
        [
            ["几千人在线", "用户打开页面或保持会话，不一定都在同时发消息。", "在线数不是最危险的指标，关键是同时触发 Agent 请求的并发数。"],
            ["几千请求并发", "同一时间大量请求进入 /agent/chat。", "这会占用后端线程，并可能打到 LLM 限流。"],
            ["系统崩溃", "不一定 JVM 退出，更常见是线程耗尽、超时、限流和用户一直转圈。", "高并发下要追求可控降级，而不是无限硬扛。"],
        ],
        [1.2, 2.55, 2.75],
    )

    doc.add_heading("2. 当前链路和主要瓶颈", level=1)
    add_para(doc, "当前 AI 客服请求链路可以概括为：")
    add_note(
        doc,
        "请求链路",
        "前端发送消息 -> /agent/chat -> 鉴权和会话识别 -> RAG 检索 pgvector -> 工具查询 MySQL/Redis -> 调用 DeepSeek/OpenAI 兼容接口 -> 保存会话记忆 -> 返回用户。",
        fill="F4F6F9",
    )
    add_matrix(
        doc,
        ["瓶颈点", "为什么危险", "可能表现"],
        [
            ["LLM 调用", "耗时最长，且受第三方 QPS、并发数、token 速率限制。", "响应慢、429 限流、超时、失败率升高。"],
            ["Tomcat 工作线程", "同步调用 LLM 时，一个请求会长时间占住一个线程。", "请求排队，普通接口也变慢。"],
            ["MySQL 连接池", "Agent 工具会查订单、购物车、用户、商品。", "连接等待，接口超时，数据库 CPU 升高。"],
            ["pgvector/RAG", "每次问题都可能向量检索和 rerank。", "PostgreSQL 延迟升高，RAG 结果返回慢。"],
            ["Redis 会话记忆", "高频读写会话和限流数据。", "Redis 延迟升高，但通常不是第一瓶颈。"],
        ],
        [1.45, 2.55, 2.5],
    )

    doc.add_heading("3. 总体设计目标", level=1)
    add_bullets(
        doc,
        [
            "保护核心业务：AI 客服高峰不能拖垮登录、商品浏览、下单、支付等电商主链路。",
            "控制外部依赖风险：DeepSeek/LLM 失败时系统要能降级，不让所有请求一直等待。",
            "控制成本：大模型调用是高成本资源，不能把几千并发原样转发给 LLM。",
            "保证用户体验：可以排队、流式返回、提示繁忙，但不能无响应。",
            "容量可验证：通过压测数据证明系统能承载多少并发，而不是凭感觉估计。",
        ],
    )

    doc.add_heading("4. 分阶段落地方案", level=1)

    doc.add_heading("阶段一：入口限流、超时、熔断", level=2)
    add_para(
        doc,
        "这是最先做的保护层。AI 接口耗时长、成本高，如果入口没有保护，几千并发会直接压垮线程池和外部 LLM API。",
    )
    add_bullets(
        doc,
        [
            "单用户限流：例如每个用户每 3 秒最多发送 1 条消息，避免重复点击和刷接口。",
            "单 IP 限流：防止恶意流量或异常客户端。",
            "全局限流：限制整个 Agent 服务同时处理的 AI 请求数，例如 100 个。",
            "超时控制：连接超时 3 秒，读取超时 15-30 秒，总超时不超过 30 秒。",
            "熔断降级：DeepSeek 连续超时或失败时，短时间内不再调用 LLM，改走 FAQ 或转人工话术。",
        ],
    )
    add_para(
        doc,
        "可选技术：Bucket4j、Resilience4j、Sentinel、Redis RateLimiter、Spring Cloud Gateway 限流。",
    )

    doc.add_heading("阶段二：异步化和流式返回", level=2)
    add_para(
        doc,
        "当前同步模式最大的问题是请求线程会一直等待 LLM 返回。更好的做法是把 Agent 请求转为异步任务，由后台 Worker 控制消费速度，并通过 SSE 或 WebSocket 推送结果。",
    )
    add_matrix(
        doc,
        ["步骤", "设计", "价值"],
        [
            ["1", "前端发送消息，后端生成 conversationId/messageId。", "建立可追踪的会话和消息状态。"],
            ["2", "请求写入 Redis Stream、RabbitMQ、Kafka 或 RocketMQ。", "削峰填谷，避免瞬时流量打穿 LLM。"],
            ["3", "Agent Worker 池从队列消费任务。", "控制真正进入 LLM 的并发数。"],
            ["4", "通过 SSE/WebSocket 流式返回 token 或最终答案。", "用户更早看到反馈，减少长 HTTP 等待。"],
        ],
        [0.75, 3.0, 2.75],
    )
    add_note(
        doc,
        "面试重点",
        "并发用户数和 LLM 并发数不能画等号。几千用户在线时，只允许有限数量的请求同时进入大模型，其余请求排队、降级或提示繁忙。",
        fill="FFF7E6",
    )

    doc.add_heading("阶段三：LLM 并发池和队列背压", level=2)
    add_para(doc, "即使异步化，也不能无限堆积任务。需要为 LLM 调用设置明确的并发池和队列容量。")
    add_bullets(
        doc,
        [
            "llm.max-concurrent：例如最多 50 或 100 个 LLM 请求同时执行。",
            "llm.queue-size：例如队列最多积压 1000 条，超过后直接返回繁忙。",
            "llm.timeout：单次 LLM 调用最长 30 秒，超时后失败并释放资源。",
            "llm.retry：只对网络抖动做有限重试，不对明确限流做无脑重试。",
            "priority queue：VIP、支付相关、售后紧急问题可以更高优先级。",
        ],
    )

    doc.add_heading("阶段四：缓存和减少不必要的 LLM 调用", level=2)
    add_para(
        doc,
        "真正能抗高并发的 AI 客服，不是所有请求都直接进大模型，而是尽量在进入 LLM 前解决一部分请求。",
    )
    add_matrix(
        doc,
        ["缓存对象", "示例", "收益"],
        [
            ["FAQ 标准问答", "退货政策、发货时间、保修多久。", "高频问题直接返回，降低 LLM 调用量。"],
            ["RAG 检索结果", "query hash -> top documents。", "减少 pgvector 压力。"],
            ["完整答案缓存", "归一化问题 -> 标准回答。", "命中后毫秒级返回。"],
            ["用户上下文", "用户资料、最近订单、购物车摘要。", "减少 MySQL 工具调用。"],
        ],
        [1.5, 2.6, 2.4],
    )

    doc.add_heading("阶段五：服务拆分和水平扩容", level=2)
    add_para(
        doc,
        "MVP 阶段可以单体，但几千并发时建议把 Agent 服务拆出来，让 AI 高峰不影响普通电商业务。",
    )
    add_note(
        doc,
        "推荐部署形态",
        "Nginx/Gateway -> backend 多实例处理普通业务；agent-service 多实例处理 AI 客服；Redis 负责会话、限流、队列和缓存；MySQL 存业务数据；PostgreSQL + pgvector 存知识库。",
        fill="F4F6F9",
    )
    add_bullets(
        doc,
        [
            "普通业务服务：登录、商品、购物车、订单、支付。",
            "Agent 服务：RAG、工具调用编排、LLM 调用、会话记忆。",
            "独立扩容：AI 高峰时只扩 Agent Worker，不必整体扩电商后端。",
            "故障隔离：LLM 故障不影响核心交易链路。",
        ],
    )

    doc.add_heading("阶段六：连接池、线程池和数据库调优", level=2)
    add_para(doc, "调优必须基于压测和监控结果，不能盲目把线程池、连接池调大。")
    add_matrix(
        doc,
        ["对象", "建议关注", "说明"],
        [
            ["Tomcat", "max threads、accept-count、请求超时。", "同步请求多时最容易被占满。"],
            ["Hikari", "maximum-pool-size、connection-timeout。", "连接池过大可能反而压垮 MySQL。"],
            ["Redis", "连接池、慢命令、内存、key 过期策略。", "用于限流、队列、会话和缓存。"],
            ["pgvector", "HNSW 索引、topK、相似度阈值、缓存。", "控制 RAG 延迟。"],
            ["JVM", "堆内存、GC、线程数、CPU。", "观察高并发下是否出现频繁 Full GC。"],
        ],
        [1.15, 2.4, 2.95],
    )

    doc.add_heading("5. 压测和验证方案", level=1)
    add_para(
        doc,
        "压测要分阶段，不是一上来就打 5000 并发。每一档压测后都要看错误率、p95/p99、队列长度、数据库连接池、LLM 限流和最终数据一致性。",
    )
    add_numbers(
        doc,
        [
            "先跑并发集成测试：库存不超卖、购物车不重复、订单状态只流转一次、默认地址最终只有一个。",
            "再用 k6/JMeter/Locust 做接口压测：从 50、100、300、500、1000、2000、5000 VU 逐步升高。",
            "最后做业务链路压测：登录 -> 商品浏览 -> 发送客服问题 -> 查询订单 -> 追问 -> 下单。",
        ],
    )
    add_matrix(
        doc,
        ["指标", "建议阈值", "含义"],
        [
            ["HTTP 错误率", "< 1%", "超过说明服务或依赖已经不稳定。"],
            ["p95 响应时间", "普通接口 < 1s；AI 首包 < 3s", "关注大多数用户体验。"],
            ["p99 响应时间", "AI 完整回答可放宽，但必须可控", "关注尾部延迟。"],
            ["队列长度", "不能持续增长", "持续增长说明消费能力不足。"],
            ["LLM 限流次数", "越低越好", "用于校准 LLM 并发池。"],
            ["数据库连接池使用率", "长期不要接近 100%", "否则请求会排队等待连接。"],
        ],
        [1.55, 1.8, 3.15],
    )

    doc.add_heading("6. 和当前 Nexus 项目的对应关系", level=1)
    add_bullets(
        doc,
        [
            "当前入口是 /agent/chat，需要加限流、超时、熔断和鉴权保护。",
            "当前会话记忆有 RedisSessionStore，但 Spring AI advisor 仍有内存窗口，需要继续演进为真正 Redis-backed chat memory 或每轮从 Redis 回灌。",
            "当前 RAG 依赖 PostgreSQL + pgvector，压测时要单独观察 pgvector 查询耗时。",
            "当前工具调用会查 MySQL 的商品、订单、购物车、用户信息，应对高频用户上下文做短时缓存。",
            "当前 AI 调用是同步的，后续应改成队列 + Worker + SSE/WebSocket 流式返回。",
        ],
    )

    doc.add_heading("7. 面试回答模板", level=1)
    add_note(
        doc,
        "可直接背诵",
        "如果几千人同时使用 AI 客服，我不会让所有请求同步直连大模型。我的方案是分层治理：第一层做限流、超时和熔断，防止流量打穿系统；第二层把 Agent 请求异步化，通过 MQ 或 Redis Stream 排队，由固定大小的 Worker 池调用 LLM；第三层用 SSE 或 WebSocket 做流式返回，避免 HTTP 长时间阻塞；第四层对 FAQ、RAG 检索结果和用户上下文做缓存，减少不必要的大模型调用；第五层将 Agent 服务从主业务服务拆出来，支持独立水平扩容；最后通过 k6/JMeter 分阶段压测，观察 p95、p99、错误率、线程池、连接池、队列长度和 LLM 限流情况，根据瓶颈调优。",
        fill="E8EEF5",
    )

    doc.add_heading("8. 面试追问与回答", level=1)
    add_matrix(
        doc,
        ["面试官追问", "推荐回答"],
        [
            [
                "为什么不能直接扩容后端？",
                "因为瓶颈不只在后端实例，还在 LLM API、RAG、数据库连接池和线程阻塞。直接扩容可能把更多请求打到外部模型，导致更快限流。",
            ],
            [
                "为什么要异步？",
                "AI 请求耗时长，同步模式会占住 Web 线程。异步队列可以削峰填谷，Worker 池可以控制真正进入 LLM 的并发。",
            ],
            [
                "如何保证用户体验？",
                "用 SSE/WebSocket 流式返回，让用户先看到首 token 或排队状态；超过容量时给明确繁忙提示，而不是一直转圈。",
            ],
            [
                "如何证明系统能扛住？",
                "通过并发集成测试证明数据正确性，通过 k6/JMeter 压测证明容量，观察 p95/p99、错误率、队列长度、连接池和 LLM 限流。",
            ],
            [
                "LLM 挂了怎么办？",
                "熔断后降级到 FAQ/RAG 标准答案，必要时转人工，同时不影响订单、支付等核心链路。",
            ],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("9. 一页总结", level=1)
    add_bullets(
        doc,
        [
            "AI 客服高并发的核心不是让所有请求都进大模型，而是控制进入大模型的请求数量。",
            "限流、超时、熔断是第一道保护；异步队列和 Worker 池是容量治理核心。",
            "SSE/WebSocket 流式返回能改善用户体验，但不能替代后端限流和队列。",
            "缓存 FAQ、RAG 结果和用户上下文，可以显著减少 LLM 调用。",
            "Agent 服务应与普通业务服务拆分，避免 AI 高峰拖垮交易链路。",
            "压测要看 p95/p99、错误率、队列长度、连接池和 LLM 限流，不只看 QPS。",
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("用于面试复习：AI 客服高并发治理，从 MVP 到可扩展架构")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("777777")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
